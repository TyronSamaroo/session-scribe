from pydantic import BaseModel
from typing import Dict, Optional
from docx import Document
import openai
import os
from dotenv import load_dotenv
load_dotenv()
import logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

openai.api_key = os.environ["OPENAI_API_KEY"]
client = openai.OpenAI(api_key=openai.api_key)

class AudioFileInfo(BaseModel):
    audio_file_path: str

class TranscriptionOutput(BaseModel):
    transcription_text: str


class MeetingMinutesOutput(BaseModel):
    abstract_summary: Optional[str]
    detailed_outline: Optional[str]
    key_points: Optional[str]
    action_items: Optional[str]
    sentiment: Optional[str]

class SaveAsDocxInput(BaseModel):
    minutes: MeetingMinutesOutput
    filename: str


def transcribe_audio(input: AudioFileInfo) -> Optional[TranscriptionOutput]:
    try:
        with open(input.audio_file_path, 'rb') as audio_file:
            logger.info("Transcribing audio file...")
            transcription = client.audio.transcriptions.create(model="whisper-1", file=audio_file)
            logger.info(transcription)
        return TranscriptionOutput(transcription_text=transcription.text)
    except Exception as e:
        logger.error(f"An error occurred while transcribing audio: {e}")
        return None

def meeting_minutes(transcription: TranscriptionOutput) -> Optional[MeetingMinutesOutput]:
    try:
        text = transcription.transcription_text
        logger.info("Extracting abstract summary...")
        abstract_summary = abstract_summary_extraction(text)
        logger.info("Generating detailed outline...")
        detailed_outline = generate_detailed_outline(text)
        logger.info("Extracting key points...")
        key_points = key_points_extraction(text)
        logger.info("Extracting action items...")
        action_items = action_item_extraction(text)
        logger.info("Performing sentiment analysis...")
        sentiment = sentiment_analysis(text)
        logger.info("Creating MeetingMinutesOutput object...")
        return MeetingMinutesOutput(
            abstract_summary=abstract_summary,
            detailed_outline=detailed_outline,
            key_points=key_points,
            action_items=action_items,
            sentiment=sentiment
        )
    except Exception as e:
        logger.error(f"An error occurred while generating meeting minutes: {e}")
        return None
    
def abstract_summary_extraction(transcription: str) -> str:
    response = client.chat.completions.create(
        model="gpt-4",
        temperature=0,
        messages=[
            {
                "role": "system",
                "content": "You are a highly skilled AI trained in language comprehension and summarization. I would like you to read the following text and summarize it into a concise abstract paragraph. Aim to retain the most important points, providing a coherent and readable summary that could help a person understand the main points of the discussion without needing to read the entire text. Please avoid unnecessary details or tangential points."
            },
            {
                "role": "user",
                "content": transcription
            }
        ]
    )
    print("Abstract summary extraction response: ", response)
    return response.choices[0].message.content


def key_points_extraction(transcription: str) -> str:
    response = client.chat.completions.create(
        model="gpt-4",
        temperature=0,
        messages=[
            {
                "role": "system",
                "content": "You are a proficient AI with a specialty in distilling information into key points. Based on the following text, identify and list the main points that were discussed or brought up. These should be the most important ideas, findings, or topics that are crucial to the essence of the discussion. Your goal is to provide a list that someone could read to quickly understand what was talked about."
            },
            {
                "role": "user",
                "content": transcription
            }
        ]
    )
    return response.choices[0].message.content
    

def action_item_extraction(transcription: str) -> str:
    response = client.chat.completions.create(
        model="gpt-4",
        temperature=0,
        messages=[
            {
                "role": "system",
                "content": "You are an AI expert in analyzing conversations and extracting action items. Please review the text and identify any tasks, assignments, or actions that were agreed upon or mentioned as needing to be done. These could be tasks assigned to specific individuals, or general actions that the group has decided to take. Please list these action items clearly and concisely."
            },
            {
                "role": "user",
                "content": transcription
            }
        ]
    )
    return response.choices[0].message.content

def generate_detailed_outline(transcription: str) -> str:
    response = client.chat.completions.create(
        model="gpt-4",
        temperature=0,
        messages=[
            {
                "role": "system",
                "content": "You are an AI expert in analyzing conversations, understanding their context, and creating detailed outlines. Please review the text, consider its context, and generate a structured outline that includes the main topics, subtopics, and key points discussed. The outline should be organized, clear, and provide a comprehensive overview of the discussion, keeping in mind the context and nuances of the conversation."
            },
            {
                "role": "user",
                "content": transcription
            }
        ]
    )
    return response.choices[0].message.content


def sentiment_analysis(transcription: str) -> str:
    response = client.chat.completions.create(
        model="gpt-4",
        temperature=0,
        messages=[
            {
                "role": "system",
                "content": "As an AI with expertise in language and emotion analysis, your task is to analyze the sentiment of the following text. Please consider the overall tone of the discussion, the emotion conveyed by the language used, and the context in which words and phrases are used. Indicate whether the sentiment is generally positive, negative, or neutral, and provide brief explanations for your analysis where possible."
            },
            {
                "role": "user",
                "content": transcription
            }
        ]
    )
    return response.choices[0].message.content

def save_as_docx(input: SaveAsDocxInput):
    doc = Document()
    for key, value in input.minutes.items():
        # Replace underscores with spaces and capitalize each word for the heading
        heading = ' '.join(word.capitalize() for word in key.split('_'))
        doc.add_heading(heading, level=1)
        doc.add_paragraph(value)
        # Add a line break between sections
        doc.add_paragraph()
    doc.save(input.filename)

def transcribe_audio_if_confirmed(input_data: AudioFileInfo) -> Optional[TranscriptionOutput]:
    print("Please note that transcribing the audio will incur charges to your OpenAI account.")
    confirm_transcription = input("Are you sure you want to transcribe the audio? (yes/no): ")
    if confirm_transcription.lower() == 'yes':
        return transcribe_audio(input_data.audio_file_path)
    else:
        print("Transcription cancelled.")
        return None

def generate_minutes_if_confirmed(transcription: TranscriptionOutput) -> Optional[MeetingMinutesOutput]:
    if transcription is None:
        return None

    print("Please note that generating meeting minutes with GPT-4 will incur charges to your OpenAI account.")
    confirm_minutes = input("Are you sure you want to generate meeting minutes? (yes/no): ")
    if confirm_minutes.lower() == 'yes':
        return meeting_minutes(transcription.transcription_text)
    else:
        print("Meeting minutes generation cancelled.")
        return None
